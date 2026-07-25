from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from datetime import datetime
import os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def h1(text):
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


doc.add_heading("FlashTap Terminal — Session Work Summary", level=0)
subtitle = doc.add_paragraph()
run = subtitle.add_run(
    "Detailed record of investigation, findings, fixes, and builds from this Cursor chat session.\n"
    f"Document generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} (local)\n"
    "Project: C:\\RN\\FlashTapTerminal\n"
    "Device under discussion: Wiseasy P5 (Finatic-UAT / production terminals)\n"
    "Latest production build from this session: v1.59 (versionCode 60)"
)
run.italic = True

h1("1. Session Scope and Constraints")
para(
    "This chat spanned two major workstreams: (A) forensic investigation of built-in receipt "
    "printing on the Wiseasy P5 using the WisePos SDK6 path, and (B) diagnosis and fix of a "
    "Sale → Charge payment UI bug that showed success or failure without launching Finatic. "
    "A production APK was also built twice (v1.58 then v1.59)."
)
h3("Standing constraints observed")
bullet("Do not modify the web/backend unless asked.")
bullet("Do not redesign print architecture or replace the WiseSdk AAR.")
bullet("Do not change Bluetooth printing or receipt rendering casually.")
bullet("Do not invent speculative printer fixes without evidence (later reinforced by user).")
bullet("Built-in printing remains off by default until hardware-proven.")
bullet("Do not commit/push unless explicitly asked.")

h1("2. Chronological Overview")
bullet(
    "Continued from a long prior thread on POS receipt printing "
    "(phases A–E, native hardening, staging APKs)."
)
bullet(
    "User directed a forensic comparison of FlashTap vs official SDK6 demo "
    "(com.smartpos.sdkdemo), with one-difference-at-a-time implementation — "
    "then reversed course to evidence-only investigation."
)
bullet(
    "Decompiled WiseSdk_P; proved 7101 comes from queryIntentServices returning "
    "zero matches for com.wisepos.aidl.service."
)
bullet(
    "Investigated why the query returns zero; then pivoted to discovering which "
    "printer Binder stacks actually exist (Finatic Device SDK, SDK4 WangPOS, "
    "SDK6 WisePos AIDL)."
)
bullet("Paused printer work; built production APK v1.58.")
bullet(
    "User reported Sale → Charge showing payment successful without Finatic; "
    "root-caused to stale AsyncStorage payment state; fixed and shipped v1.59."
)
bullet(
    "User asked why testing device did not show the bug; explained state-dependent nature."
)
bullet(
    "User showed Payment FAILED appearing before Process Payment on an older APK; "
    "confirmed same persistence class of bug (FAILED leftover)."
)

h1("3. Built-in Printer Investigation (Wiseasy P5 / SDK6)")

h2("3.1 Problem statement")
para(
    "FlashTap’s built-in printer path uses WiseSdk_P_1.29_00a_24041501.aar and "
    "WisePosSdk.initPosSdk(), binding to the Android intent action "
    "com.wisepos.aidl.service. On the P5 under test, printing failed. Device probe "
    "and error UI reported matches=0 for that action (model=P5, targetSdk=29), and "
    "initPosSdk failed with error 7101 (ERR_SDK_INVALID_PARAMETER)."
)

h2("3.2 Confirmed shared SDK")
bullet("AAR: WiseSdk_P_1.29_00a_24041501.aar")
bullet(
    "SHA-256 identical between FlashTap and SDK6 demo package: "
    "946516FF12B5E799C904469C52EE8B1FE532AC30F901AC057091303C7A062E93"
)
bullet(
    "Same WisePosSdk / Printer API / com.wisepos.aidl.service client strings "
    "embedded in both APKs"
)
bullet("Also ships WiseSdk_D (device/scanner) alongside P in FlashTap")

h2("3.3 Early mistaken comparisons (corrected in-session)")
para(
    "Prior work had mixed evidence from apps that print on the same hardware but "
    "use different stacks. This session reinforced those corrections:"
)

h3("Finatic Cashier (com.wiseasy.finatic.cashier)")
bullet("Does NOT use the public WisePos AIDL path as its primary print stack.")
bullet(
    "Uses com.wiseasy.devicesdk.printer.* "
    "(PrintManager, LocalPrinter, PrintIntentService)."
)
bullet(
    "Optional system shared library: wiseasy.device.sdk "
    "(uses-library, required=false)."
)
bullet(
    "Signed by Wiseasy corporate certificate — different from FlashTap and SDK Demo."
)
bullet(
    "May embed WisePos / libbasebinder strings as alternate backends, but declared "
    "built-in path is Device SDK."
)

h3('User-provided “SDK Demo” APK that printed (9ca68225…)')
bullet(
    "Package: com.wpos.sdkdemo (SDK4 / WangPOS family), NOT com.smartpos.sdkdemo."
)
bullet("Uses wangpos.sdk4.libbasebinder.Printer and com.pos.permission.PRINTER.")
bullet(
    "Binds via wangpos.sdk4.*.service.BinderPoolService (base / emv / keymanager)."
)
bullet(
    "Native libs: libbasebinder, libemvbinder, libkeymanagerbinder."
)
bullet("No WisePosSdk / com.wisepos.aidl.service in that APK.")

h3("Official SDK6 demo (correct apples-to-apples reference)")
bullet("Package: com.smartpos.sdkdemo / SDKDemo_1.31")
bullet(
    "MainActivity.onCreate: WisePosSdk.initPosSdk(this) then "
    "WiseDeviceSdk.initDeviceSdk(this)"
)
bullet(
    "PrinterActivity: getPrinter → initPrinter → setGrayLevel → status → font → "
    "addSingleText → startPrinting → feedPaper"
)
bullet(
    "targetSdk 29; no QUERY_ALL_PACKAGES; no printer <queries>; Android Debug signer"
)

h2("3.4 Decompilation of WisePosSdk.initPosSdk (CFR)")
para("Critical client logic proven from WiseSdk_P classes.jar:")
bullet('Builds Intent("com.wisepos.aidl.service").')
bullet(
    "getExplicitIntent: PackageManager.queryIntentServices(intent, flags=0) "
    "must return exactly 1 match."
)
bullet("0 or 2+ matches → null → onInitPosFail(7101).")
bullet(
    "On match: context.getApplicationContext().bindService(explicitIntent, …, "
    "BIND_AUTO_CREATE)."
)
bullet(
    "On connect: IServiceManager.register(null, new Binder()) then onInitPosSuccess()."
)
bullet(
    "getPrinter() returns PrinterImpl stub even when unbound; initPrinter() returns "
    "7102 if IServiceManager is null."
)
bullet(
    "Client SDK does NOT check app package name, signing certificate, process name, "
    "or task affinity for the resolve step."
)
bullet(
    "Only one WisePos action string exists: com.wisepos.aidl.service "
    "(no fallback actions)."
)
bullet(
    "WiseDeviceSdk separately uses com.wisedevice.aidl.service "
    "(not required for printer)."
)

h2("3.5 FlashTap probe vs SDK lookup")
para(
    "FlashTap’s probeUsdkService was compared to WisePosSdk.getExplicitIntent. "
    "For the path that produces 7101, they are functionally identical: same action, "
    "flags=0, queryIntentServices. The probe also tries MATCH_ALL for diagnostics only. "
    "Probe false-negative vs SDK was ruled out with high confidence."
)

h2("3.6 Why queryIntentServices returned zero (evidence-based)")
para(
    "On-device screenshot evidence: matches=0, model=P5, targetSdk=29 for "
    "com.wisepos.aidl.service."
)
bullet(
    "Package visibility filtering (Android 11+ for targetSdk ≥ 30) was RULED OUT "
    "because FlashTap APK and on-device error both showed targetSdk 29."
)
bullet(
    "Exact proven statement: for FlashTap’s UID/user, PackageManager had no "
    "resolvable service matching that action."
)
bullet(
    "Sub-cause not fully separable without adb dumpsys: host APK absent vs not "
    "exported vs disabled vs other user."
)
bullet(
    "No host USDK APK was found in the SDK6 client distribution on disk; provider "
    "package name remains unidentified from artifacts alone."
)
bullet(
    "adb devices was empty during investigation — live device enumeration of "
    "installed services was blocked."
)

h2("3.7 Three printer stacks on this firmware family")
para(
    "Ranked discovery of what “actually exists” for printing "
    "(from APK forensics, not live dumpsys):"
)

h3("Stack A — SDK6 WisePos AIDL (FlashTap + official SDK6 demo)")
bullet("Action: com.wisepos.aidl.service")
bullet("API: WisePosSdk → PrinterImpl → AIDL IPrinter")
bullet("On this P5: not resolvable (matches=0) for FlashTap")
bullet(
    "FlashTap and official SDK6 demo use the SAME stack — FlashTap is not "
    "“wrong” relative to SDK6 client code"
)

h3("Stack B — Finatic Device SDK")
bullet(
    "com.wiseasy.devicesdk.printer.* + optional wiseasy.device.sdk shared library"
)
bullet("In-app PrintIntentService")
bullet("Not an intent-filter replacement for com.wisepos.aidl.service")

h3("Stack C — SDK4 WangPOS")
bullet("wangpos.sdk4.libbasebinder.Printer via BinderPoolService")
bullet("Permission: com.pos.permission.PRINTER")
bullet("Different Binder family entirely")

h2('3.8 Conclusion on “has the service been replaced?”')
bullet(
    "No evidence of a renamed drop-in replacement for com.wisepos.aidl.service "
    "with the same WisePos client API."
)
bullet(
    "Working print on this unit/family is explained by parallel stacks "
    "(Device SDK / WangPOS), while the SDK6 AIDL host is not visible to FlashTap."
)
bullet(
    "No code change recommended for printer until a live host package/service is "
    "confirmed on a connected P5."
)

h2("3.9 Code changes related to printer (limited)")
para(
    "Early in the session (before the user switched to evidence-only mode), "
    "FlashTap’s WisePosSdkBootstrap init order was swapped to match the demo: "
    "initPosSdk then initDeviceSdk. Later investigation concluded init-order "
    "cosmetics cannot fix matches=0. User then asked to stop speculative "
    "demo-matching changes. No further printer architecture changes were made "
    "in the later evidence-only phase."
)
bullet(
    "Relevant files: WisePosSdkBootstrap.kt, WiseSdk6PrinterModule.kt, "
    "MainActivity.kt, AndroidManifest.xml (<queries>, QUERY_ALL_PACKAGES), "
    "android/build.gradle (targetSdk 29)."
)
bullet(
    "These had largely been introduced in prior conversation turns; this session "
    "focused on forensics."
)

h1("4. Sale → Charge Payment Bug (Finatic Not Launched)")

h2("4.1 User-reported symptom")
para(
    "When creating an order through Sale → View Cart → Charge, the terminal showed "
    "“Payment successful” without bringing up the Finatic / WiseCashier payment app. "
    "Dashboard Order History showed mixed COMPLETED and PENDING card orders "
    "(e.g. #393 COMPLETED earlier; #396 PENDING), consistent with UI lying about "
    "payment while the backend order remained unpaid."
)

h2("4.2 Root cause")
para(
    "The payment state machine (PaymentStateMachine / AsyncStorage key "
    "PAYMENT_STATE_STORAGE_KEY) previously persisted PAYMENT_SUCCESS "
    "(and PAYMENT_FAILED) across navigations. On the next Sale → Charge, "
    "PaymentScreen hydrated that stale state and immediately showed success "
    "(or failure) without calling Finatic."
)
bullet(
    "SUCCESS leftover → “Payment successful” without Finatic (user’s first report)."
)
bullet(
    "FAILED leftover → “Payment failed / Payment was cancelled or failed” before "
    "Process Payment (user’s later screenshot on an older APK)."
)
bullet("Same mechanism; opposite stored state.")
bullet(
    "Why testing device often looked fine: fresh installs, cleared data, fewer "
    "Charge→pay→Charge loops, debug rebuilds wiping storage. Production devices "
    "retain AsyncStorage across many sales."
)

h2("4.3 Related code paths reviewed")
bullet(
    "POSCartScreen.handleCharge → createPOSOrder → navigation.replace(Payment)."
)
bullet(
    "PaymentScreen.handleProcessPayment → processPaymentIntent → "
    "PaymentModule.launchPayment (intent com.wiseasy.transaction.call) → "
    "MainActivity.onActivityResult → completePayment."
)
bullet(
    "MainActivity only resolves success on result \"00\" with a non-blank "
    "transactionID (voucher)."
)
bullet(
    "processPaymentIntent was hardened to refuse success without voucherNo "
    "(no invented FT-* references)."
)

h2("4.4 Fixes shipped in v1.59")
bullet("Never persist SUCCESS/FAILED — only crash-recover PAYMENT_IN_PROGRESS.")
bullet("Drop legacy PAYMENT_SUCCESS on hydrate; clear state when orderId mismatches.")
bullet(
    "clearPersistedPaymentState() called from POSCartScreen before navigating "
    "to Payment after Charge."
)
bullet(
    "PaymentScreen waits for isHydrated before rendering success/fail UI "
    "(avoids flash of stale state)."
)
bullet(
    "processPaymentIntent requires a real Finatic voucher/transaction ID for success."
)
bullet("Version bumped to 1.59 / versionCode 60 (JS APP_VERSION and Gradle).")

h2("4.5 Files touched for payment fix")
bullet("src/components/PaymentStateMachine.tsx")
bullet("src/screens/POSCartScreen.tsx")
bullet("src/screens/PaymentScreen.tsx")
bullet("src/lib/payment.ts (voucher requirement — confirmed/hardened)")
bullet("android/app/build.gradle")
bullet("src/constants/index.ts")

h1("5. APK Builds Produced This Session")

h2("5.1 Production v1.58")
bullet('Requested after pausing printer work (“leave receipt issue for tomorrow”).')
bullet(
    "Path: android/app/build/outputs/apk/production/release/"
    "app-production-release.apk"
)
bullet("Package: com.flashtap.pos")
bullet("versionName 1.58 / versionCode 59")
bullet(
    "Note: this build may not have included the full payment-persistence fix "
    "depending on working-tree state at build time; v1.59 is the payment-fix build."
)

h2("5.2 Production v1.59")
bullet("Built after payment bug fix.")
bullet("Same output path (overwrites production release APK).")
bullet("versionName 1.59 / versionCode 60")
bullet(
    r"Full path: C:\RN\FlashTapTerminal\android\app\build\outputs\apk\production\release\app-production-release.apk"
)
bullet("Approximate size: ~70.9 MB")

h1("6. Open Items / Deferred Work")
bullet(
    "Built-in P5 printing via WisePos AIDL remains unresolved pending on-device "
    "adb enumeration of which host packages/services exist (wiseasy.device.sdk, "
    "wangpos BinderPoolService, USDK for com.wisepos.aidl.service)."
)
bullet(
    "Confirm official SDK6 demo (com.smartpos.sdkdemo) initPosSdk success/fail "
    "on the same P5 where FlashTap reports matches=0."
)
bullet(
    "If Device SDK or WangPOS is the only viable stack on this firmware, that "
    "would be a deliberate architecture decision — not pursued in this session "
    "per “no speculative redesign” rules."
)
bullet(
    "Payment fix should be verified on the field device after installing v1.59: "
    "Sale → Charge should show clean Process Payment, then Finatic on tap."
)
bullet(
    "Git: work was largely uncommitted relative to user “don’t commit unless "
    "asked”; branch context included pos-receipt-email-v144 / later checkpoints."
)

h1("7. Key Evidence Artifacts Referenced")
bullet("SDK6 zip / SmartPosSdkDemo source and SDKDemo_1.31 APK")
bullet(
    r"Finatic Cashier APK: Downloads\38aefc1df4ae3cb23974b1a75d1bf366.apk"
)
bullet(
    r"SDK4 demo APK: Downloads\9ca682251ebdd27f8b80ccb7fdab2b03.apk (com.wpos.sdkdemo)"
)
bullet(r"Decompiled WiseSdk_P under %TEMP%\forensic_wisesdk_p\cfr_out")
bullet("On-device Settings error screenshot: matches=0, model=P5, targetSdk=29")
bullet(
    "Dashboard Order History screenshot (24/07/2026) showing COMPLETED #393 "
    "and PENDING #396"
)
bullet(
    "Payment screen screenshot showing FAILED before Process Payment on "
    "pre-1.59 build"
)

h1("8. Ranked Technical Conclusions")

h2("Printer / USDK")
bullet(
    "1. FlashTap faithfully queries com.wisepos.aidl.service and gets zero "
    "matches on the tested P5 (High confidence)."
)
bullet(
    "2. That empty result is not explained by a FlashTap probe bug or by "
    "targetSdk package-visibility (targetSdk 29) (High confidence)."
)
bullet(
    "3. Apps that print on this firmware family use other stacks "
    "(Device SDK / WangPOS), not proof that WisePos AIDL is present "
    "(High confidence)."
)
bullet(
    "4. Exact host package for WisePos AIDL on this image is unproven without "
    "adb (Medium — blocked by no device connected during forensics)."
)

h2("Payment UI")
bullet(
    "1. Instant success/fail without Finatic was caused by persisted payment "
    "machine state in AsyncStorage (High confidence)."
)
bullet(
    "2. v1.59 addresses persistence, Charge-time clear, hydrate gating, and "
    "voucher requirement (High confidence for the reported UI bug)."
)
bullet(
    "3. Field verification on the production terminal after install remains "
    "the final confirmation (Operational)."
)

h1("9. Appendix — Important Constants and Identifiers")
bullet("WisePos USDK action: com.wisepos.aidl.service")
bullet("WiseDevice action: com.wisedevice.aidl.service")
bullet("Finatic/WiseCashier payment action: com.wiseasy.transaction.call")
bullet("FlashTap payment appId (native): wz66363c6bb9592fb5")
bullet("Error 7101: ERR_SDK_INVALID_PARAMETER (cannot resolve USDK intent)")
bullet(
    "Error 7102: ERR_SDK_SERVICE_NOT_CONNECTED (printer APIs without live bind)"
)
bullet("Production applicationId: com.flashtap.pos")
bullet(
    "Staging applicationId (from earlier APK dumps): com.flashtap.pos.staging"
)

para("")
footer = doc.add_paragraph()
fr = footer.add_run(
    "End of summary. This document describes work performed in the Cursor agent "
    "chat session and related forensic analysis on local APKs/sources. It is not "
    "a formal vendor certification report."
)
fr.italic = True
fr.font.size = Pt(9)

out_path = os.path.join(
    r"C:\RN\FlashTapTerminal", "FlashTap_Session_Summary_2026-07-25.docx"
)
doc.save(out_path)
print(out_path)
print(os.path.getsize(out_path))
